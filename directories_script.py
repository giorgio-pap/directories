#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Aug 13 12:13:43 2021

@author: giorgiopapitto

Name: Directories 
A non creative writing project

"""
    
import os
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

#create poems
#define function that returns the list of all files and subdirectories
#within a root directory and subdirectories
def listdirs(rootdir):
    for file in os.listdir(rootdir):
        d = os.path.join(rootdir, file)
        if os.path.isdir(d):
            folders_list.append(d)
            listdirs(d)

rootdir = "/Users/giorgiopapitto/"
folders_list = []
n = 0 

try:
    listdirs(rootdir)
except OSError:
    pass 

try:
    for folder in folders_list:
        filelist = os.listdir(folder)
        if 2 < len(filelist) < 100: #get only directories with a specific number of files
            n = n+1
            n_str=str(n)
            #create an empy text file
            text_poem = open(rootdir + "Desktop/Directories/Poems/Poem" + n_str + ".txt","w+")
            #only get what comes after the last / in the directory name
            folder.rsplit('/',1)
            folder = folder.rsplit('/',1)[1]
            #write poem title
            text_poem.write(folder + "\n\n\n")
            
            #put every file in the folder as a line
            for line in filelist:
                text_poem.writelines(line + "\n") 
        else:
            continue

except OSError:
    pass 

#delete last poem (usually empty)
os.remove(rootdir + "Desktop/Directories/Poems/Poem" + n_str + ".txt") 

#create a list containing all the poems
list_poems = []
folder = rootdir + "Desktop/Directories/Poems/"
for poem in os.listdir(folder):
    if poem.endswith('.txt'):
        list_poems.append(poem)

print(len(list_poems))
#create 5 volumes of 95 poems

#sort list numerically        
from natsort import natsorted
list_poems = natsorted(list_poems)

start = 0
end = 95

for x in range(1,6): #because of 5 volumes
    poems_titles = [] #this will be used for getting the titles
    volume_files = list_poems[start:end] #get poems for each volume
    string = rootdir + "Desktop/Directories/Poems/"
    volume_files = [string + y for y in volume_files]
    start = start + 95
    end = end + 95
    #create doc file
    doc = Document()
    doc.add_heading("Volume " + str(x), 0)
    doc.add_page_break()

    for poem_a in volume_files:
            
        with open(poem_a) as f:
            lines = f.readlines()
            first = True

      
            for line in lines:
                line = line.replace("\n","")
                
                #if first line, use it as a title
                if first:
                    doc.add_heading(line,1)
                    first = False
                    poems_titles.append(line)
                
                #otherwise as line of the poem    
                else:
                    doc_para = doc.add_paragraph(line)
                    paragraph_format = doc_para.paragraph_format
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
           
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Garamond'
                    font.size = Pt(12) 
        
            doc.add_page_break()
                
    doc.save(rootdir + "Desktop/Directories/Volumes/Volume_" + str(x) + ".docx")